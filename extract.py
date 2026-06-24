"""
VOUAR Brain — Motor de Extração de Conteúdo
Lê todos os documentos-chave e gera JSONs estruturados para o dashboard.
"""

import os, json, re
from pathlib import Path
import docx

BASE = Path("/Users/andreluismribeiro/VouAR - MTG")
OUT  = Path("/Users/andreluismribeiro/vouar-brain/data/inventario")
OUT.mkdir(exist_ok=True)

DOCUMENTOS = [
    # Manuscritos
    {"id": "livro-mtg",   "tipo": "livro",   "metodo": "MTG",   "path": BASE / "01_LIVROS/MTG/MTX_MTG_Livro_COMPLETO_v2.1.docx"},
    {"id": "livro-focus", "tipo": "livro",   "metodo": "FOCUS", "path": BASE / "01_LIVROS/FOCUS/MTX_FOCUS_Livro_COMPLETO_v1.3.docx"},
    {"id": "livro-mte",   "tipo": "livro",   "metodo": "MTE",   "path": BASE / "01_LIVROS/MTE/MTX_MTE_Livro_REVISADO_v2_12.docx"},
    {"id": "livro-mtc",   "tipo": "livro",   "metodo": "MTC",   "path": BASE / "01_LIVROS/MTC/MTX_MTC_Livro_COMPLETO_v1.3.docx"},
    # PRDs e specs
    {"id": "civ-prd",     "tipo": "prd",     "metodo": "CIV",   "path": BASE / "04_PRODUTOS/CIV_PRD_v6.0.docx"},
    {"id": "mtg-brain",   "tipo": "spec",    "metodo": "MTG",   "path": BASE / "04_PRODUTOS/2026-06-17_MTG_Brain_Inteligencia_Negocio_v5.docx"},
    # Instrumentos
    {"id": "ferramentas", "tipo": "kit",     "metodo": "MTG",   "path": BASE / "09_INSTRUMENTOS/Ferramentas.docx"},
    {"id": "mapa-eco",    "tipo": "mapa",    "metodo": "VOUAR", "path": BASE / "09_INSTRUMENTOS/Mapa_do_Ecossistema_Comercial_MTG_v2_0.docx"},
]

SEP_RE     = re.compile(r'^[━─—\-=]{10,}')
PARTE_RE   = re.compile(r'^(PARTE|CAPÍTULO|MÓDULO|SEÇÃO|BLOCO|ETAPA|STEP)\s+\w', re.I)
CAPS_RE    = re.compile(r'^[A-ZÁÉÍÓÚÀÃÕÂÊÔ\s\d\.\/\-]{6,}$')
CONCEPT_RE = re.compile(r'\b(framework|método|metodologia|modelo|sistema|ferramenta|canvas|template|planner|score|checklist|mapa|processo|protocolo|POP|SCI|RACI)\b', re.I)
DECISION_RE= re.compile(r'\b(decidimos|definimos|optamos|escolhemos|resolvemos|a decisão|ficou definido|ficou acordado|não inclui|removemos|foi removido|perdemos|falta|pendente|a fazer)\b', re.I)


def extract_paragraphs(path):
    try:
        d = docx.Document(str(path))
        return [p.text.strip() for p in d.paragraphs if p.text.strip()]
    except Exception as e:
        return []


def detect_sections(paras):
    """Identifica seções a partir de padrões de texto (sem estilos heading)."""
    sections = []
    current = {"titulo": "INÍCIO", "conteudo": [], "nivel": 1}
    for p in paras:
        if SEP_RE.match(p):
            continue
        if PARTE_RE.match(p) or (CAPS_RE.match(p) and len(p) < 80 and not p.startswith('"')):
            if current["conteudo"]:
                sections.append(current)
            nivel = 1 if PARTE_RE.match(p) else 2
            current = {"titulo": p, "conteudo": [], "nivel": nivel}
        else:
            current["conteudo"].append(p)
    if current["conteudo"]:
        sections.append(current)
    return sections


def extract_concepts(paras):
    found = set()
    for p in paras:
        for m in CONCEPT_RE.finditer(p):
            # pega a frase ao redor
            start = max(0, m.start() - 40)
            end   = min(len(p), m.end() + 60)
            snippet = p[start:end].strip()
            found.add(snippet)
    return sorted(found)[:30]


def extract_decisions(paras):
    found = []
    for p in paras:
        if DECISION_RE.search(p) and len(p) > 30:
            found.append(p[:200])
    return found[:20]


def extract_bold_terms(path):
    """Extrai termos em negrito — geralmente conceitos-chave."""
    try:
        d = docx.Document(str(path))
        bold = set()
        for p in d.paragraphs:
            for run in p.runs:
                if run.bold and run.text.strip() and len(run.text.strip()) > 3:
                    bold.add(run.text.strip()[:80])
        return sorted(bold)[:50]
    except:
        return []


def process_doc(doc_meta):
    path = doc_meta["path"]
    if not path.exists():
        return {"id": doc_meta["id"], "erro": f"Arquivo não encontrado: {path.name}"}

    print(f"  Processando {doc_meta['id']}...")
    paras     = extract_paragraphs(path)
    sections  = detect_sections(paras)
    concepts  = extract_concepts(paras)
    decisions = extract_decisions(paras)
    bold      = extract_bold_terms(path)

    result = {
        "id":         doc_meta["id"],
        "tipo":       doc_meta["tipo"],
        "metodo":     doc_meta["metodo"],
        "arquivo":    path.name,
        "total_paras": len(paras),
        "secoes":     [{"titulo": s["titulo"], "nivel": s["nivel"], "n_paras": len(s["conteudo"])} for s in sections],
        "conceitos":  concepts,
        "decisoes_encontradas": decisions,
        "termos_em_negrito": bold,
        "titulo_detectado": paras[0] if paras else "",
    }
    return result


def main():
    print("VOUAR Brain — Extração de Conteúdo")
    print("=" * 50)
    todos = []
    for doc in DOCUMENTOS:
        r = process_doc(doc)
        todos.append(r)
        out_path = OUT / f"{doc['id']}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(r, f, ensure_ascii=False, indent=2)
        if "erro" in r:
            print(f"  ✗ {doc['id']}: {r['erro']}")
        else:
            print(f"  ✓ {doc['id']}: {len(r['secoes'])} seções · {len(r['termos_em_negrito'])} termos em negrito")

    # índice geral
    with open(OUT / "_indice.json", "w", encoding="utf-8") as f:
        json.dump({"documentos": todos}, f, ensure_ascii=False, indent=2)

    print("\nConcluído. Arquivos em:", OUT)


if __name__ == "__main__":
    main()
